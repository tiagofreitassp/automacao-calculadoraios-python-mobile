import os, shutil
from time import sleep
from appium import webdriver
from appium.options.common import AppiumOptions
from docx import Document
from docx.shared import Inches

dir = '../evidencias/'

xcodeOrgId = "3P3JYXL37U"
xcodeSigningId = "iPhone 15"
udid = "A4EAE15D-1507-40A9-8B4D-96F007852E68"
platformName = "iOS"
deviceName = "iPhone 13"
platformVersion = "17.0"
bundleId = "com.mouredev.iOS-Calculator"
automationName = "XCUITest"
driverApp = ""

class mobileDriver():
    def criarPastaEvidencia(self, nPasta):
            d = nPasta
            if os.path.exists(d) == True:
                print("Diretório já existe")
                #os.rmdir(d)
                shutil.rmtree(d)
                print('Diretório apagado')
                os.makedirs(d)
                print('Diretório recriado')
            else:
                print("Diretório não existe")
                os.makedirs(d)
                print('Diretório criado')

    def gerarScreenshot(self, nPasta, nEvidencia):
        sleep(1)
        self.driver.get_screenshot_as_file(nPasta + "/" + nEvidencia + ".png")

    def criarDocumentoDeEvidencia(self, diretorioEvidencia,id ,nomeEvidencia):
        try:
            document = Document()

            document.add_heading('Evidências: Calculadora Android', 0)
            p = document.add_paragraph(nomeEvidencia)

            document.add_paragraph(id+'_Tela01')
            document.add_picture(diretorioEvidencia + '/Ev1.png', width=Inches(4.14))
            #document.add_page_break()

            document.add_paragraph(id+'_Tela02')
            document.add_picture(diretorioEvidencia + '/Ev2.png', width=Inches(4.14))
            #document.add_page_break()

            document.add_paragraph(id + '_Tela03')
            document.add_picture(diretorioEvidencia + '/Ev3.png', width=Inches(4.14))
            # document.add_page_break()

            document.save(diretorioEvidencia + '/' + nomeEvidencia + '.docx')
            print("Documento com as evidencias gerada com sucesso!")
        except:
            print("Não foi possivel criar o documento com as evidencias!")

    def criarDriver(self):
        desired_caps = {
            "platformName": platformName,
            "deviceName": deviceName,
            "automationName": automationName,
            "xcodeSigningId": xcodeSigningId,
            "udid": udid,
            "platformVersion": platformVersion,
            "bundleId": bundleId,
            # "xcodeOrgId": xcodeOrgId,
            # desired_caps['app'] = driverApp
        }

        option = AppiumOptions().load_capabilities(desired_caps)
        self.driver = webdriver.Remote(command_executor=f"http://127.0.0.1:4723", options=option)
        self.driver.implicitly_wait(12)

    def fecharDriver(self):
        self.driver.quit()